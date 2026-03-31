import os, zipfile
from datetime import datetime
from typing import List, Dict

class FileExporter:
    def __init__(self, output_dir='output'):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def _bname(self, repo):
        ts = datetime.now().strftime('%Y%m%d_%H%M%S')
        safe = repo.replace('/','_').replace(' ','_')
        return f'repo_analysis_{safe}_{ts}'

    def export_txt(self, text, repo_name):
        p = os.path.join(self.output_dir, f'{self._bname(repo_name)}.txt')
        with open(p, 'w', encoding='utf-8') as f: f.write(text)
        print(f'  TXT: {p}')
        return p

    def export_pdf(self, text, repo_name):
        from fpdf import FPDF
        p = os.path.join(self.output_dir, f'{self._bname(repo_name)}.pdf')
        class P(FPDF):
            def header(self):
                self.set_font('Helvetica','B',10)
                self.cell(0,10,f'Analysis: {repo_name}',0,1,'C')
                self.ln(3)
            def footer(self):
                self.set_y(-15)
                self.set_font('Helvetica','I',8)
                self.cell(0,10,f'Page {self.page_no()}',0,0,'C')
        pdf = P()
        pdf.set_auto_page_break(True, 15)
        pdf.add_page()
        pdf.set_font('Helvetica','',7)
        for line in text.split(chr(10)):
            try:
                safe = line.encode('latin-1',errors='replace').decode('latin-1')
                pdf.multi_cell(0, 3, safe, 0, 'L')
            except: continue
        pdf.output(p)
        print(f'  PDF: {p}')
        return p

    def export_word(self, text, repo_name, scan_data=None):
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        p = os.path.join(self.output_dir, f'{self._bname(repo_name)}.docx')
        doc = Document()
        t = doc.add_heading('Repository Analysis Report', 0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub.add_run(repo_name)
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0,102,204)
        doc.add_page_break()
        in_code = False
        code_buf = []
        for line in text.split(chr(10)):
            if line.startswith('```'):
                if in_code:
                    code = chr(10).join(code_buf)
                    if code.strip():
                        cp = doc.add_paragraph()
                        r = cp.add_run(code)
                        r.font.name = 'Courier New'
                        r.font.size = Pt(7)
                        cp.paragraph_format.left_indent = Inches(0.3)
                    code_buf = []
                    in_code = False
                else: in_code = True
                continue
            if in_code: code_buf.append(line); continue
            s = line.strip()
            if not s: doc.add_paragraph(); continue
            if s.startswith('='*20): continue
            pr = doc.add_paragraph()
            r = pr.add_run(s)
            r.font.size = Pt(9)
        doc.save(p)
        print(f'  Word: {p}')
        return p

    def export_zip(self, files, repo_name):
        p = os.path.join(self.output_dir, f'{self._bname(repo_name)}.zip')
        with zipfile.ZipFile(p, 'w', zipfile.ZIP_DEFLATED) as zf:
            for f in files:
                if os.path.exists(f): zf.write(f, os.path.basename(f))
        print(f'  ZIP: {p}')
        return p

    def export_all(self, report_text, repo_name, scan_data=None, formats=None, create_zip=False):
        if formats is None: formats = ['txt']
        exported, paths = {}, []
        for fmt in formats:
            try:
                if fmt=='txt': pa = self.export_txt(report_text, repo_name)
                elif fmt=='pdf': pa = self.export_pdf(report_text, repo_name)
                elif fmt=='word': pa = self.export_word(report_text, repo_name, scan_data)
                else: continue
                exported[fmt] = pa; paths.append(pa)
            except Exception as e: print(f'  Error {fmt}: {e}')
        if create_zip and paths:
            try: exported['zip'] = self.export_zip(paths, repo_name)
            except Exception as e: print(f'  ZIP error: {e}')
        return exported